# -*- coding: utf-8 -*-
"""
build_hfc_docx.py
Rebuilds HFC_ScaleUp_Report.docx using python-docx, matching the original
PDF formatting and adding Section 4.3 (Yang-Cussler comparison).
"""

import io, math, os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np

# ── Output path ────────────────────────────────────────────────────────────────
OUT = os.path.join(os.path.dirname(__file__), 'HFC_ScaleUp_Report.docx')

# ── Color palette ──────────────────────────────────────────────────────────────
NAVY      = RGBColor(0x1A, 0x3A, 0x5C)   # section headings
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)   # body text
MID_GRAY  = RGBColor(0x55, 0x55, 0x55)
TBL_HDR   = RGBColor(0x1A, 0x3A, 0x5C)   # table header fill (navy)
TBL_ALT   = RGBColor(0xE8, 0xF0, 0xF8)   # alternating row fill
BOX_BG    = RGBColor(0xE8, 0xF5, 0xE9)   # green callout bg
BOX_BDR   = RGBColor(0x2E, 0x7D, 0x32)   # green callout border
INFO_BG   = RGBColor(0xE3, 0xF2, 0xFD)   # blue callout bg
INFO_BDR  = RGBColor(0x15, 0x65, 0xC0)   # blue callout border
RULE_CLR  = RGBColor(0x1A, 0x3A, 0x5C)   # horizontal rule colour
EQ_BG     = RGBColor(0xF5, 0xF5, 0xF5)   # equation box background

# ── Font names ─────────────────────────────────────────────────────────────────
BODY_FONT  = 'Calibri'
HEAD_FONT  = 'Calibri'
MONO_FONT  = 'Courier New'

# ══════════════════════════════════════════════════════════════════════════════
#  LOW-LEVEL XML HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def _rgb_hex(rgb: RGBColor) -> str:
    return str(rgb)  # python-docx RGBColor.__str__ returns 'RRGGBB'

def _set_cell_color(cell, rgb: RGBColor):
    """Fill a table cell with a solid colour."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  _rgb_hex(rgb))
    tcPr.append(shd)

def _set_cell_borders(cell, top=None, bottom=None, left=None, right=None,
                      color='1A3A5C', sz='6'):
    """Set individual cell borders."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom),
                      ('left', left), ('right', right)]:
        if val:
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'),   val)
            b.set(qn('w:sz'),    sz)
            b.set(qn('w:color'), color)
            tcBorders.append(b)
    tcPr.append(tcBorders)

def _no_borders(cell):
    """Remove all borders from a cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'bottom', 'left', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),   'none')
        b.set(qn('w:sz'),    '0')
        b.set(qn('w:color'), 'auto')
        tcBorders.append(b)
    tcPr.append(tcBorders)

def _cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom),
                      ('left', left), ('right', right)]:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'),    str(val))
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)

def _set_row_height(row, height_pt):
    trPr = row._tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_pt * 20)))
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)

def _keep_with_next(para):
    pPr = para._p.get_or_add_pPr()
    kn  = OxmlElement('w:keepNext')
    pPr.append(kn)

def _add_left_border(cell, color='1A3A5C', sz='24'):
    """Add a thick left border to a cell (used for callout boxes)."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'bottom', 'right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),   'none')
        b.set(qn('w:sz'),    '0')
        b.set(qn('w:color'), 'auto')
        tcBorders.append(b)
    b = OxmlElement('w:left')
    b.set(qn('w:val'),   'single')
    b.set(qn('w:sz'),    sz)
    b.set(qn('w:color'), color)
    tcBorders.append(b)
    tcPr.append(tcBorders)

def _set_table_borders(table, color='1A3A5C', sz='4'):
    tbl  = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'bottom', 'left', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),   'single')
        b.set(qn('w:sz'),    sz)
        b.set(qn('w:color'), color)
        tblBorders.append(b)
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)

def _set_page_margins(doc, top=1.0, bottom=1.0, left=1.1, right=1.1):
    section = doc.sections[0]
    section.top_margin    = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin   = Inches(left)
    section.right_margin  = Inches(right)

def _para_spacing(para, before=0, after=4, line=None):
    pPr  = para._p.get_or_add_pPr()
    pSpc = OxmlElement('w:spacing')
    pSpc.set(qn('w:before'), str(before * 20))
    pSpc.set(qn('w:after'),  str(after  * 20))
    if line is not None:
        pSpc.set(qn('w:line'),     str(line * 240))
        pSpc.set(qn('w:lineRule'), 'auto')
    pPr.append(pSpc)

# ══════════════════════════════════════════════════════════════════════════════
#  HIGH-LEVEL STYLE HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def _run(para, text, bold=False, italic=False, size=11, color=None, font=None,
         underline=False, mono=False):
    run = para.add_run(text)
    run.bold      = bold
    run.italic    = italic
    run.underline = underline
    run.font.name = (MONO_FONT if mono else (font or BODY_FONT))
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run

def body_para(doc, text='', bold=False, italic=False, size=11,
              color=None, align=WD_ALIGN_PARAGRAPH.LEFT,
              before=0, after=6, indent=0):
    p = doc.add_paragraph()
    p.alignment = align
    _para_spacing(p, before=before, after=after, line=1.15)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    if text:
        _run(p, text, bold=bold, italic=italic, size=size,
             color=color or DARK_GRAY)
    return p

def section_heading(doc, text, level=1, before=14, after=4):
    """Numbered section heading in navy bold."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _para_spacing(p, before=before, after=after)
    sz = {1: 13, 2: 12, 3: 11}.get(level, 11)
    _run(p, text, bold=True, size=sz, color=NAVY, font=HEAD_FONT)
    _keep_with_next(p)
    return p

def horizontal_rule(doc, before=8, after=8):
    p = doc.add_paragraph()
    _para_spacing(p, before=before, after=after)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:color'), _rgb_hex(RULE_CLR))
    pBdr.append(bot)
    pPr.append(pBdr)
    return p

def bullet_para(doc, text, size=11, indent=0.25, before=2, after=2):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _para_spacing(p, before=before, after=after, line=1.15)
    p.paragraph_format.left_indent  = Inches(indent + 0.2)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    _run(p, text, size=size, color=DARK_GRAY)
    return p

def equation_box(doc, text, before=4, after=4):
    """Equation displayed in a shaded single-cell table."""
    tbl  = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    _set_cell_color(cell, EQ_BG)
    _no_borders(cell)
    _cell_margins(cell, top=60, bottom=60, left=120, right=120)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _para_spacing(p, before=0, after=0)
    _run(p, text, italic=True, size=10.5, color=DARK_GRAY, mono=False)
    spacer = doc.add_paragraph()
    _para_spacing(spacer, before=before, after=after)
    return tbl

def callout_box(doc, title, body_lines, bg_color=INFO_BG, border_color=INFO_BDR,
                title_color=None, before=10, after=10):
    """Single-cell shaded callout box with bold title."""
    tbl  = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.rows[0].cells[0]
    _set_cell_color(cell, bg_color)
    _add_left_border(cell, color=_rgb_hex(border_color), sz='24')
    _cell_margins(cell, top=80, bottom=80, left=160, right=120)
    # Title
    tp = cell.paragraphs[0]
    tp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _para_spacing(tp, before=0, after=4)
    _run(tp, title, bold=True, size=11,
         color=title_color or RGBColor.from_string(_rgb_hex(border_color)))
    # Body lines
    for line in body_lines:
        bp = cell.add_paragraph()
        bp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _para_spacing(bp, before=0, after=2, line=1.15)
        _run(bp, line, size=10.5, color=DARK_GRAY)
    spacer = doc.add_paragraph()
    _para_spacing(spacer, before=0, after=before)
    return tbl

def data_table(doc, headers, rows, col_widths=None, caption=None,
               alt_rows=True, before=6, after=6):
    """
    Standard data table with navy header and optional alternating row shading.
    col_widths: list of Inches values (optional)
    """
    ncols = len(headers)
    tbl   = doc.add_table(rows=1 + len(rows), cols=ncols)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_borders(tbl, color='1A3A5C', sz='4')

    # Set column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = w

    # Header row
    hdr = tbl.rows[0]
    for j, h in enumerate(headers):
        cell = hdr.cells[j]
        _set_cell_color(cell, TBL_HDR)
        _cell_margins(cell, top=60, bottom=60, left=90, right=60)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _para_spacing(p, before=0, after=0)
        _run(p, h, bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF))

    # Data rows
    for i, row_data in enumerate(rows):
        row = tbl.rows[i + 1]
        bg  = TBL_ALT if (alt_rows and i % 2 == 1) else RGBColor(0xFF, 0xFF, 0xFF)
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            _set_cell_color(cell, bg)
            _cell_margins(cell, top=50, bottom=50, left=90, right=60)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _para_spacing(p, before=0, after=0)
            _run(p, str(cell_text), size=9.5, color=DARK_GRAY)

    if caption:
        cp = doc.add_paragraph()
        _para_spacing(cp, before=4, after=10)
        _run(cp, caption, italic=True, size=9.5, color=MID_GRAY)
    else:
        sp = doc.add_paragraph()
        _para_spacing(sp, before=0, after=before)

    return tbl

def figure_para(doc, caption_text, before=4, after=12):
    p = doc.add_paragraph()
    _para_spacing(p, before=before, after=after)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, caption_text, italic=True, size=9.5, color=MID_GRAY)
    return p

def title_block(doc):
    """Cover title + field details table + abstract."""
    # Main title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(p, before=0, after=4, line=1.2)
    _run(p, 'Hollow-Fiber Membrane Contactor Scale-Up Design\n'
            'for Citric Acid Recovery from Agricultural Waste Fermentation Broth',
         bold=True, size=16, color=NAVY)

    # Subtitle
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(p2, before=0, after=14, line=1.1)
    _run(p2, 'Feasibility Study and Module Sizing Based on the Yang-Cussler Model',
         bold=False, italic=True, size=12, color=MID_GRAY)

    # Field details table
    meta = [
        ('Author',         'Ogaga Maxwell Okedi, M.S. Chemical Engineering, Florida A&M University'),
        ('Date',           'April 2025'),
        ('Classification', 'Supporting Technical Document, EB-2 NIW Petition (Prong II, Section 3.4)'),
        ('Project',        'BioForge \u2014 Integrated DOE + ML Waste-to-Value Optimization Platform (v0.1)'),
        ('Repository',     'https://github.com/ogaga-ai/BioForge'),
    ]
    tbl = doc.add_table(rows=len(meta), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_borders(tbl, color='1A3A5C', sz='4')
    col_w = [Inches(1.3), Inches(5.4)]
    for i, (k, v) in enumerate(meta):
        bg = TBL_ALT if i % 2 == 1 else RGBColor(0xFF, 0xFF, 0xFF)
        for j, txt in enumerate([k, v]):
            cell = tbl.rows[i].cells[j]
            cell.width = col_w[j]
            _set_cell_color(cell, bg)
            _cell_margins(cell, top=55, bottom=55, left=90, right=60)
            p = cell.paragraphs[0]
            _para_spacing(p, before=0, after=0)
            _run(p, txt, bold=(j == 0), size=10, color=DARK_GRAY)

    doc.add_paragraph()

    # Abstract heading
    ah = doc.add_paragraph()
    _para_spacing(ah, before=10, after=4)
    _run(ah, 'Abstract', bold=True, size=12, color=NAVY)

    # Abstract body
    body_para(doc,
        'This report presents a scale-up feasibility study for a hollow-fiber membrane contactor (HFC) '
        'module designed to recover citric acid from yam peel agricultural waste fermentation broth, '
        'building on the foundational design equations of Yang and Cussler (1986). This work grows '
        'directly out of a rigorous technical project I completed during graduate training, in which I '
        'studied hollow-fiber membrane contactors from first principles \u2014 tracing the physics of '
        'interfacial mass transfer through the Sieder-Tate Sherwood-number correlation and the '
        'three-resistance framework, then adapting and validating those equations as a practical tool '
        'for scale-up design. This report is the engineering realization of that work, applying the '
        'validated framework to citric acid recovery at a BioForge-predicted yield of 43.08\u202fg/L '
        '(Okedi et al., 2024). Pilot-scale calculations yield K\u2097 = 1.1324\u00d710\u207b\u00b3 cm/s, '
        'with 5 modules in series achieving 95.4% recovery, exceeding the 95% design target and '
        'outperforming the conventional extraction baseline of 85%. Industrial scale-up to 1,000\u202fL '
        'requires 50 modules (37.7\u202fm\u00b2 total area). This study positions the HFC module as the '
        'Stage\u202f2 separation layer within BioForge.',
        before=0, after=10)

    horizontal_rule(doc)

# ══════════════════════════════════════════════════════════════════════════════
#  FIGURE GENERATORS
# ══════════════════════════════════════════════════════════════════════════════

def _fig_sh_gz():
    fig, ax = plt.subplots(figsize=(5.5, 3.5), dpi=150)
    Gz_arr = np.linspace(1, 120, 300)
    Sh_arr = 1.86 * Gz_arr ** (1/3)
    ax.plot(Gz_arr, Sh_arr, 'b-', lw=2, label='Sh = 1.86 Gz\u00b9\u141f\u00b3 (Sieder-Tate)')
    ax.plot(42.10, 6.47, 'ro', ms=9, zorder=5,
            label='Design point: Gz=42.1, Sh=6.47')
    ax.annotate('BioForge\nStage 2\ndesign point\n(Gz=42.1, Sh=6.47)',
                xy=(42.10, 6.47), xytext=(60, 4.8),
                arrowprops=dict(arrowstyle='->', color='red', lw=1.4),
                fontsize=8, color='darkred')
    ax.set_xlabel('Graetz Number (Gz)', fontsize=10)
    ax.set_ylabel('Sherwood Number (Sh)', fontsize=10)
    ax.set_title('Sh\u2013Gz Correlation: Sieder-Tate (Yang-Cussler, 1986)', fontsize=10)
    ax.legend(fontsize=8)
    ax.grid(True, alpha=0.3)
    ax.set_xlim(0, 120); ax.set_ylim(0, 10)
    plt.tight_layout()
    buf = io.BytesIO(); fig.savefig(buf, format='png', bbox_inches='tight'); plt.close(fig)
    buf.seek(0); return buf

def _fig_recovery():
    fig, ax = plt.subplots(figsize=(5.5, 3.5), dpi=150)
    n_mods = np.arange(1, 9)
    eta_s  = 0.459
    eta_t  = 1 - (1 - eta_s) ** n_mods * 100
    eta_t  = (1 - (1 - eta_s) ** n_mods) * 100
    cols   = ['#90CAF9' if n != 5 else '#388E3C' for n in n_mods]
    bars   = ax.bar(n_mods, eta_t, color=cols, edgecolor='#1A3A5C', lw=0.8)
    ax.axhline(95, color='red', ls='--', lw=1.5, label='95% target')
    ax.axhline(85, color='orange', ls=':', lw=1.5, label='Conventional baseline (85%)')
    for bar, val in zip(bars, eta_t):
        ax.text(bar.get_x() + bar.get_width()/2, val + 0.5,
                f'{val:.1f}%', ha='center', va='bottom', fontsize=8)
    ax.set_xlabel('Number of Modules in Series', fontsize=10)
    ax.set_ylabel('Cumulative Recovery (%)', fontsize=10)
    ax.set_title('Cumulative Citric Acid Recovery vs Modules in Series', fontsize=10)
    ax.set_xticks(n_mods)
    ax.set_ylim(0, 105)
    ax.legend(fontsize=8)
    ax.grid(True, axis='y', alpha=0.3)
    plt.tight_layout()
    buf = io.BytesIO(); fig.savefig(buf, format='png', bbox_inches='tight'); plt.close(fig)
    buf.seek(0); return buf

def _fig_schematic():
    fig, ax = plt.subplots(figsize=(6.0, 3.2), dpi=150)
    ax.set_xlim(0, 10); ax.set_ylim(0, 4); ax.axis('off')
    colors = ['#1565C0','#1976D2','#1E88E5','#42A5F5','#90CAF9']
    recov  = [45.9, 70.8, 84.1, 90.5, 95.4]
    for i in range(5):
        x = 0.5 + i * 1.8
        rect = mpatches.FancyBboxPatch((x, 1.0), 1.3, 1.8,
            boxstyle='round,pad=0.05', linewidth=1.5,
            edgecolor='#0D47A1', facecolor=colors[i])
        ax.add_patch(rect)
        ax.text(x + 0.65, 1.92, f'M{i+1}', ha='center', va='center',
                color='white', fontsize=10, fontweight='bold')
        ax.text(x + 0.65, 1.25, f'{recov[i]:.1f}%', ha='center', va='center',
                color='white', fontsize=8)
        if i < 4:
            ax.annotate('', xy=(x + 1.8, 1.92), xytext=(x + 1.3, 1.92),
                        arrowprops=dict(arrowstyle='->', color='#1A3A5C', lw=1.5))
    ax.text(0.0, 1.92, 'Feed\n43.08 g/L', ha='center', va='center',
            fontsize=8, color='#1A3A5C', fontweight='bold')
    ax.text(9.6, 1.92, 'Raffinate', ha='center', va='center',
            fontsize=8, color='#1A3A5C', fontweight='bold')
    ax.text(5.0, 3.5,
            'BioForge Stage 2 Pilot-Scale HFC Array  |  5 \u00d7 30 cm modules  |  Q\u2097 = 13.9 cm\u00b3/s',
            ha='center', va='center', fontsize=9, color='#1A3A5C', style='italic')
    ax.text(5.0, 0.4,
            'Alamine 336 / n-heptane (shell side, counter-current)',
            ha='center', va='center', fontsize=8.5, color='#555555')
    plt.tight_layout()
    buf = io.BytesIO(); fig.savefig(buf, format='png', bbox_inches='tight'); plt.close(fig)
    buf.seek(0); return buf

# ══════════════════════════════════════════════════════════════════════════════
#  DOCUMENT SECTIONS
# ══════════════════════════════════════════════════════════════════════════════

def section_1(doc):
    section_heading(doc, '1.  Introduction', level=1)
    body_para(doc,
        'Industrial separation processes account for 40% to 70% of total plant operating costs in '
        'U.S. chemical and biochemical manufacturing. As the United States transitions toward '
        'bio-based manufacturing of organic acids from agricultural waste, the efficiency of the '
        'downstream separation step becomes a critical bottleneck determining commercial viability.')
    body_para(doc,
        'Conventional liquid-liquid extraction systems for organic acid recovery from fermentation '
        'broth face well-documented limitations: flooding and channeling in packed columns, solvent '
        'emulsification contaminating the product stream, high energy costs associated with solvent '
        'recovery, and large equipment footprints. These limitations are especially pronounced in '
        'small-to-medium agricultural waste biorefineries, where compact and modular separation '
        'equipment is essential.')
    body_para(doc,
        'Hollow-fiber membrane contactors (HFCs) offer a compelling alternative. Formalized by '
        'Yang and Cussler (1986) for gas-liquid and liquid-liquid extraction, HFCs enable mass '
        'transfer between immiscible phases without dispersing one phase within the other, '
        'eliminating flooding, emulsification, and channeling while offering specific surface areas '
        'of 500 to 5,000\u202fm\u00b2/m\u00b3, far exceeding the 100 to 200\u202fm\u00b2/m\u00b3 '
        'typical of packed columns.')
    body_para(doc,
        'This report addresses a specific engineering design question central to the BioForge '
        'Stage\u202f2 development roadmap: what hollow-fiber module configuration is required to '
        'recover citric acid from fermentation broth produced when BioForge-optimized conditions '
        'are applied to agricultural waste? This question sits at the intersection of two bodies of '
        'work I know deeply. During my graduate training, I undertook a project focused on the '
        'technical study of hollow-fiber membrane contactors. Hollow fiber membrane-based systems '
        'are used in bioprocessing to selectively separate and recover target chemicals across a '
        'stable, non-dispersive interface. That study centered on the design equations of Yang and '
        'Cussler (AIChE Journal, 1986), which I analyzed, extended, and validated as a design tool '
        'for real-scale applications. This report is the direct engineering outcome of that '
        'training: it takes the validated Yang-Cussler framework and applies it to the specific '
        'separation challenge that BioForge Stage\u202f2 must solve: sizing the module, confirming '
        'recovery, and establishing a clear path from pilot to industrial scale.')

    callout_box(doc,
        'Connection to BioForge Platform',
        [
            'BioForge Stage 1 uses Design of Experiments (DOE) and machine learning to identify '
            'optimal fermentation conditions that maximize citric acid yield from agricultural waste. '
            'Stage 2 receives those conditions as inputs and sizes the downstream separation module '
            'required to extract the product at commercially relevant purity and recovery. This '
            'report designs that Stage 2 module.'
        ],
        bg_color=INFO_BG, border_color=INFO_BDR)

    horizontal_rule(doc)


def section_2(doc):
    section_heading(doc, '2.  Process Description: BioForge Citric Acid Production', level=1)
    section_heading(doc, '2.1  Published Process Context', level=2, before=8)
    body_para(doc,
        'The process described in this report is based on my peer-reviewed research published in '
        'Industrial Crops and Products (IF\u202f6.2) in 2024 (Okedi et al., 2024). In that study, '
        'a Box-Behnken experimental design identified optimal concentrations of three process '
        'stimulants (EDTA, coconut oil, and sodium fluoride) for citric acid production by '
        'Aspergillus niger under solid-state fermentation on yam peel substrate (Dioscorea spp.). '
        'The published experimental results, validated by an ANFIS model, yielded a maximum citric '
        'acid concentration of 43.08\u202fg/L, a 49.1% improvement over the unstimulated baseline '
        'yield of 28.90\u202fg/L. BioForge Stage\u202f1 reproduces this result computationally, '
        'achieving R\u00b2\u202f=\u202f0.99883 on the same 17-run dataset.')

    section_heading(doc, '2.2  Process Specifications', level=2, before=8)
    data_table(doc,
        headers=['Parameter', 'Value', 'Source'],
        rows=[
            ['Substrate',                   'Yam peel waste (Dioscorea spp.)',                'Okedi et al. (2024)'],
            ['Organism',                    'Aspergillus niger, solid-state fermentation',    'Okedi et al. (2024)'],
            ['Fermentation temperature',    '30\u202f\u00b0C',                                'Okedi et al. (2024)'],
            ['Fermentation duration',       '6 days',                                         'Okedi et al. (2024)'],
            ['Maximum citric acid yield',   '43.08 g/L (+49.1% vs baseline)',                 'Okedi et al. (2024)'],
            ['BioForge ML model R\u00b2',   '0.99883 (ANN, tanh, MFFF architecture)',         'BioForge Stage 1'],
            ['Optimal EDTA',                '0.30 g/L',                                       'Okedi et al. (2024) / BioForge'],
            ['Optimal Coconut Oil',         '2.50 %w/w',                                      'Okedi et al. (2024) / BioForge'],
            ['Optimal Sodium Fluoride',     '0.05 g/L',                                       'Okedi et al. (2024) / BioForge'],
            ['Feed concentration for HFC design', '43.08 g/L citric acid',                   'This study'],
        ],
        col_widths=[Inches(2.4), Inches(2.4), Inches(2.0)])

    section_heading(doc, '2.3  The Downstream Separation Bottleneck', level=2, before=8)
    body_para(doc,
        'After fermentation, the broth contains citric acid distributed among fungal biomass, '
        'residual substrate, secreted enzymes, and other metabolites. Conventional downstream '
        'processing involves precipitation with calcium hydroxide, acidulation with sulfuric acid, '
        'decolorization, and crystallization, generating calcium sulfate waste and achieving only '
        '85 to 90% recovery. Liquid-liquid extraction using reactive solvents (tertiary amines such '
        'as Alamine\u202f336) offers a more direct route, and hollow-fiber contactors are ideally '
        'suited for implementing this step, maintaining a stable aqueous-organic interface without '
        'emulsification, sized using the Yang-Cussler design equations.')
    horizontal_rule(doc)


def section_3(doc):
    section_heading(doc, '3.  Hollow-Fiber Membrane Contactor: Theoretical Foundation', level=1)
    section_heading(doc, '3.1  Module Geometries', level=2, before=8)
    section_heading(doc, '3.1.1  Cylindrical Shell Module (Selected for This Design)', level=3, before=6)
    body_para(doc,
        'The cylindrical shell module consists of microporous hollow fibers bundled inside a '
        'cylindrical shell, analogous to a shell-and-tube heat exchanger. Liquid flows through '
        'the fiber lumens and extracts solvent through the shell side. This geometry is selected '
        'for the BioForge Stage\u202f2 design because it produces well-defined lumen flow with '
        'predictable mass transfer governed by the Sieder-Tate correlation, and it is the most '
        'widely validated configuration in literature. Yang and Cussler reported flow rates of '
        '2 to 20\u202fcm\u00b3/s for this module type.')
    section_heading(doc, '3.1.2  Crossflow Module (Future Consideration)', level=3, before=6)
    body_para(doc,
        'In the crossflow configuration, liquid flows perpendicular to the fiber bundle, creating '
        'wakes that enhance shell-side mixing. Yang and Cussler reported 40% higher mass transfer '
        'coefficients for 750-fiber crossflow modules vs. 72-fiber modules. For the citric acid / '
        'Alamine\u202f336 system studied here, lumen-side resistance dominates (K \u2248 K\u2097, '
        'confirmed in Section\u202f3.3), so the parallel-flow cylindrical geometry is selected. '
        'Crossflow configurations will be evaluated in BioForge Stage\u202f2 software as an '
        'optimization option.')

    section_heading(doc, '3.2  Mass Transfer Theory', level=2, before=8)
    body_para(doc,
        'Film theory (Cussler, 1984) governs the rate of citric acid transfer from the fermentation '
        'broth into the extracting solvent across the hollow-fiber membrane. Assuming concentration '
        'changes occur only within thin films adjacent to phase boundaries:')
    equation_box(doc, 'N = K\u2097 (C\u1d35 \u2212 C\u2082\u1d64\u2097\u1d4f)    [Eq. 1]')
    body_para(doc,
        'where N = molar flux [mol/s\u00b7cm\u00b2], K = overall mass transfer coefficient [cm/s], '
        'C\u1d35 = solute concentration at the interface, C\u2082\u1d64\u2097\u1d4f = bulk '
        'concentration. The overall resistance is the sum of three resistances in series:')
    equation_box(doc, '1/K = 1/K\u2097 + 1/(K\u2098\u00b7H) + 1/(K\u1d65\u00b7H)    [Eq. 2]')
    body_para(doc,
        'where K\u2097 = liquid-phase coefficient, K\u2098 = membrane coefficient, K\u1d65 = '
        'organic-phase coefficient, H = distribution coefficient. For reactive extraction of '
        'citric acid with Alamine\u202f336, H is large (5 to 20 at pH\u202f2\u20133), making the '
        'membrane and organic-phase resistance terms negligible:')
    equation_box(doc, 'K \u2248 K\u2097    (when H >> 1, reactive extraction regime)    [Eq. 3]')

    section_heading(doc, '3.3  Yang-Cussler Design Equation', level=2, before=8)
    body_para(doc,
        'For liquid flowing through the fiber lumen past a reactive extracting solvent on the '
        'shell side, the fractional recovery is:')
    equation_box(doc, '\u03b7 = 1 \u2212 exp(\u2212K\u2097 \u00b7 A / Q\u2097)    [Eq. 4]')

    section_heading(doc, '3.4  Sieder-Tate Correlation for K\u2097', level=2, before=8)
    body_para(doc,
        'The liquid-phase mass transfer coefficient under laminar lumen-side flow is predicted '
        'by the Sieder-Tate correlation adapted for mass transfer (Yang-Cussler, 1986):')
    equation_box(doc, 'Sh = 1.86 \u00b7 Gz\u00b9\u141f\u00b3    [Eq. 5]')
    equation_box(doc, 'where:  Sh = K\u2097\u00b7d\u1d35/D        Gz = Re\u00b7Sc\u00b7(d\u1d35/L)    [Eq. 6]')
    equation_box(doc, 'Re = \u03c1\u00b7V\u2097\u00b7d\u1d35/\u03bc        Sc = \u03bc/(\u03c1\u00b7D)    [Eq. 7]')
    equation_box(doc, 'Rearranging:  K\u2097 = (1.86 \u00b7 Gz\u00b9\u141f\u00b3 \u00b7 D) / d\u1d35    [Eq. 8]')
    body_para(doc,
        'This correlation applies when Re < 2,100 (laminar regime) and was experimentally '
        'confirmed by Yang and Cussler for oxygen extraction from water across cylindrical '
        'parallel-flow modules using the same polypropylene fiber geometry selected for this design.')
    horizontal_rule(doc)


def section_4(doc):
    section_heading(doc, '4.  Experimental Validation Basis', level=1)
    section_heading(doc, '4.1  Yang and Cussler (1986): Gas-Liquid Baseline', level=2, before=8)
    body_para(doc,
        'Yang and Cussler validated the Sieder-Tate correlation (Eq.\u202f5) for three module '
        'configurations and two solute-solvent systems. Key findings directly supporting this design:')
    bullet_para(doc,
        'Experimental Sherwood numbers agreed with Sh\u202f=\u202f1.86\u00b7Gz\u00b9\u141f\u00b3 '
        'across 16, 120, and 2,100-fiber modules, confirming reliability across a range of '
        'N values.')
    bullet_para(doc,
        'Liquid-phase resistance-controlled mass transfer (K\u2097 dominant), identical to the '
        'condition established for the citric acid / Alamine\u202f336 system (Section\u202f3.3).')
    bullet_para(doc,
        'Fiber geometry (d\u1d35\u202f=\u202f0.04\u202fcm, t\u1d42\u1d43\u1d38\u1d38\u202f=\u202f'
        '0.003\u202fcm, polypropylene) is identical to the geometry used in this design, so the '
        'validated correlation applies directly.')

    section_heading(doc, '4.2  Paul and Callahan (1987): Liquid-Liquid Extension', level=2, before=8)
    body_para(doc,
        'Paul and Callahan extended Yang and Cussler\u2019s model to liquid-liquid extraction '
        '(gold from HCl solution into organic solvent). Their design equation:')
    equation_box(doc, 'Sh = 1.62 \u00b7 (V\u2097\u00b7d\u1d35\u00b2 / D\u00b7L)\u00b9\u141f\u00b3    [Eq. 9]')
    body_para(doc,
        'is structurally identical to Yang and Cussler\u2019s Eq.\u202f5, with the prefactor '
        'difference (1.62 vs 1.86) attributable to boundary condition differences. Paul and '
        'Callahan confirmed that for systems with large H, K\u202f\u2248\u202fK\u2097, validating '
        'the simplification used in this report. Their experiments demonstrated >95% extraction '
        'efficiency in multi-pass configurations, directly validating the multi-module series '
        'approach proposed in Section\u202f5.4.')
    body_para(doc,
        'The citric acid / Alamine\u202f336 system is chemically analogous to the gold / DGDE '
        'system: both involve a reactive organic extractant forming a strong complex, producing '
        'large H and making K\u2097 the sole limiting resistance. The Yang-Cussler correlation, '
        'therefore, applies to the citric acid case with high confidence.')

    # ── Section 4.3  NEW ──────────────────────────────────────────────────────
    section_heading(doc, '4.3  Extension to Citric Acid Recovery: Design Conditions vs. Yang-Cussler (1986)',
                    level=2, before=12)
    body_para(doc,
        'Yang and Cussler developed and validated their hollow-fiber contactor design equations '
        'using oxygen extracted from distilled water with a nitrogen sweep gas \u2014 a controlled, '
        'single-solute gas-liquid system chosen deliberately to isolate liquid-phase mass transfer '
        'resistance. Their best-fit experimental Sherwood number correlation for the 16-fiber '
        'cylindrical module (PS16) was:')
    equation_box(doc,
        'Sh = 1.64 \u00b7 Gz\u00b9\u141f\u00b3    '
        '(Yang-Cussler, 1986, experimental fit, PS16 module)')
    body_para(doc,
        'This fitted coefficient of 1.64 falls 12% below the Sieder-Tate theoretical prediction '
        'of 1.86 \u2014 a difference Yang and Cussler attributed to non-ideal fiber spacing and '
        'entry-length effects in their short laboratory modules (L\u202f=\u202f97.5\u202fcm). '
        'They reported mass transfer coefficients for oxygen from water in the range of '
        'approximately 2.0\u202f\u00d7\u202f10\u207b\u00b3\u202fcm/s under their test conditions, '
        'as reproduced in the MATLAB sample calculation documented in their original study.')
    body_para(doc,
        'This design applies the same Sieder-Tate framework to a fundamentally different '
        'separation problem: the recovery of citric acid from yam peel agricultural waste '
        'fermentation broth. The process data are not derived from Yang and Cussler\u2019s '
        'experimental conditions. They are sourced independently from two bodies of work:')
    bullet_para(doc,
        'The fermentation yield of 43.08\u202fg/L citric acid, established experimentally by '
        'Okedi et al.\u202f(2024) in Industrial Crops and Products using Aspergillus niger under '
        'BioForge-optimized conditions (EDTA, coconut oil, sodium fluoride on yam peel substrate) '
        'and confirmed computationally by BioForge Stage\u202f1 (R\u00b2\u202f=\u202f0.99883).')
    bullet_para(doc,
        'Physical property data for the fermentation broth '
        '(D\u202f=\u202f7.0\u202f\u00d7\u202f10\u207b\u2076\u202fcm\u00b2/s, '
        '\u03bc\u202f=\u202f0.012\u202fg/cm\u00b7s, '
        '\u03c1\u202f=\u202f1.02\u202fg/cm\u00b3) sourced from the fermentation engineering '
        'literature for citric acid aqueous solutions at 30\u202f\u00b0C and broth-level viscosity '
        '\u2014 not from the oxygen-water system used by Yang and Cussler.')
    body_para(doc,
        'Table\u202f4 below provides a direct comparison between Yang and Cussler\u2019s original '
        'experimental conditions and the design conditions applied in this study.',
        before=4, after=6)

    data_table(doc,
        headers=['Parameter', 'Yang-Cussler (1986)', 'This Design (Okedi, 2024)'],
        rows=[
            ['Target solute',
             'Oxygen (O\u2082)',
             'Citric acid from yam peel fermentation broth'],
            ['Feed source',
             'Distilled water, O\u2082-saturated',
             'Aspergillus niger fermentation broth, 43.08\u202fg/L (Okedi et al., 2024)'],
            ['Extracting phase',
             'Nitrogen sweep gas (shell side)',
             'Alamine\u202f336 (20% v/v) in n-heptane (shell side)'],
            ['Solute diffusivity, D',
             '\u223c2.0\u202f\u00d7\u202f10\u207b\u2075\u202fcm\u00b2/s (O\u2082 in water)',
             '7.0\u202f\u00d7\u202f10\u207b\u2076\u202fcm\u00b2/s (citric acid in broth)'],
            ['Fibers per module (lumen-side flow)',
             '16 to 120 fibers (PS16, PM120)',
             '2,000 fibers'],
            ['Module length',
             '97.5\u202fcm',
             '30\u202fcm'],
            ['Volumetric flow rate',
             '2 to 20\u202fcm\u00b3/s (tested range)',
             '13.89\u202fcm\u00b3/s (100\u202fL batch, 2\u202fh)'],
            ['Experimental Sh correlation',
             'Sh = 1.64\u00b7Gz\u00b9\u141f\u00b3 (PS16, fitted)',
             'Sh = 1.86\u00b7Gz\u00b9\u141f\u00b3 (Sieder-Tate, applied to new system)'],
            ['Graetz number at design point',
             'Not reported for individual modules',
             'Gz\u202f=\u202f42.10 (Re\u202f=\u202f18.8, laminar)'],
            ['K\u2097 computed',
             '\u223c2.0\u202f\u00d7\u202f10\u207b\u00b3\u202fcm/s (O\u2082 from water, MATLAB)',
             '1.1324\u202f\u00d7\u202f10\u207b\u00b3\u202fcm/s (citric acid from broth)'],
            ['Single-pass recovery',
             'Not computed',
             '45.9% per module'],
            ['Cumulative recovery',
             'Not addressed',
             '95.4% across 5 modules in series'],
            ['Industrial scale-up',
             'Not addressed',
             '50 modules, 37.7\u202fm\u00b2 total membrane area'],
        ],
        col_widths=[Inches(2.0), Inches(2.2), Inches(2.6)],
        caption='Table 4. Comparison of Yang-Cussler (1986) experimental conditions vs. '
                'this design (Okedi, 2024).')

    body_para(doc,
        'Three engineering extensions distinguish this work from Yang and Cussler\u2019s original study:',
        before=4, after=4)
    bullet_para(doc,
        'New separation system. Yang and Cussler tested oxygen and carbon dioxide removal from '
        'distilled water. Citric acid recovery from fermentation broth required independently '
        'sourced physical property data (diffusivity, viscosity, density) for an aqueous organic '
        'acid system at broth-level conditions, and a reactive extractant (Alamine\u202f336) '
        'that does not appear in Yang and Cussler\u2019s experimental program.')
    bullet_para(doc,
        'Larger module scale. Yang and Cussler\u2019s lumen-side flow experiments used 16 to '
        '120 fibers per module. This design specifies 2,000 fibers \u2014 a 17\u00d7 increase '
        'over the largest lumen-flow module they tested \u2014 while maintaining '
        'Re\u202f=\u202f18.8 (laminar) and Gz\u202f=\u202f42.10 within the validated range of '
        'the Sieder-Tate correlation.')
    bullet_para(doc,
        'Recovery quantification and multi-module staging. Yang and Cussler reported mass '
        'transfer coefficients only. This work extends the framework to compute single-pass '
        'recovery (45.9%), the minimum number of series modules required to meet a 95% design '
        'target (5 modules, 95.4% cumulative), and the full industrial scale-up configuration '
        '(50 modules, 37.7\u202fm\u00b2) \u2014 none of which appear in the 1986 paper.')

    callout_box(doc,
        'Connection to BioForge Stage 1 Yield',
        [
            'The inlet feed concentration of 43.08\u202fg/L citric acid used in all design '
            'calculations above is the experimentally verified maximum yield from Okedi et al. '
            '(2024), produced under BioForge-optimized stimulant conditions '
            '(EDTA\u202f=\u202f0.30\u202fg/L, coconut oil\u202f=\u202f2.50\u202f%w/w, '
            'NaF\u202f=\u202f0.05\u202fg/L on yam peel substrate). BioForge Stage\u202f1 '
            'confirms this yield computationally with R\u00b2\u202f=\u202f0.99883. This value '
            'is not an assumed or estimated input \u2014 it is derived directly from published, '
            'peer-reviewed experimental data and independently reproduced by the BioForge model.'
        ],
        bg_color=BOX_BG, border_color=BOX_BDR,
        title_color=RGBColor(0x1B, 0x5E, 0x20))

    horizontal_rule(doc)


def section_5(doc, fig1_buf, fig2_buf, fig3_buf):
    section_heading(doc, '5.  Module Design and Sizing: Pilot Scale', level=1)
    section_heading(doc, '5.1  Design Basis and Process Specifications', level=2, before=8)
    body_para(doc,
        'The pilot-scale design targets a 100-L batch of filtered fermentation broth processed '
        'over a 2-hour window, consistent with a typical academic or small-industry pilot '
        'fermenter. The target recovery is 95% or greater. The extracting solvent is 20%\u202fv/v '
        'Alamine\u202f336 (tri-n-octylamine) dissolved in n-heptane, with H\u202f=\u202f5 to 15 '
        'at pH\u202f2\u20133 (reactive regime).')

    data_table(doc,
        headers=['Parameter', 'Symbol', 'Value', 'Units'],
        rows=[
            ['Batch volume (pilot)',   'V\u2099\u2090\u209c\u1d9c\u02b0',       '100',      'L'],
            ['Processing time',        't',                                        '2',        'h'],
            ['Liquid flow rate',       'Q\u2097',                                  '13.89',    'cm\u00b3/s'],
            ['Feed concentration',     'C\u1d35\u207f',                            '43.08',    'g/L citric acid'],
            ['Target recovery',        '\u03b7\u209c\u2090\u02b3\u1d4f\u1d49\u209c', '95',   '%'],
            ['Fiber inner diameter',   'd\u1d35',                                  '0.04',     'cm'],
            ['Fiber outer diameter',   'd\u2092',                                  '0.046',    'cm'],
            ['Wall thickness',         't\u1d42',                                  '0.003',    'cm'],
            ['Membrane porosity',      '\u03b5',                                   '33',       '% (pore size 300\u202f\u00c5)'],
            ['Module length',          'L',                                        '30',       'cm'],
            ['Fibers per module',      'N',                                        '2,000',    '\u2014'],
            ['Membrane material',      '\u2014',                                   'Polypropylene (hydrophobic)', '\u2014'],
            ['Extracting solvent',     '\u2014',                                   '20% Alamine\u202f336 in n-heptane', '\u2014'],
            ['Operating temperature',  'T',                                        '30',       '\u00b0C'],
            ['Citric acid diffusivity','D',                                        '7.0\u202f\u00d7\u202f10\u207b\u2076', 'cm\u00b2/s'],
            ['Broth viscosity',        '\u03bc',                                   '0.012',    'g/(cm\u00b7s) [1.2\u202fcP]'],
            ['Broth density',          '\u03c1',                                   '1.02',     'g/cm\u00b3'],
        ],
        col_widths=[Inches(2.5), Inches(0.8), Inches(2.0), Inches(1.5)])

    section_heading(doc, '5.2  Step-by-Step Mass Transfer Coefficient Calculation', level=2, before=8)
    body_para(doc, 'Step 1: Cross-Sectional Area and Liquid Velocity', bold=True, before=4, after=2)
    equation_box(doc, 'A\u2091\u02b3\u02b3\u02b3\u02b3 = N\u00b7\u03c0\u00b7(d\u1d35/2)\u00b2 = 2000\u00b7\u03c0\u00b7(0.02)\u00b2 = 2.5133\u202fcm\u00b2    [Eq. 10]')
    equation_box(doc, 'V\u2097 = Q\u2097 / A\u2091\u02b3 = 13.89 / 2.5133 = 5.53\u202fcm/s    [Eq. 11]')
    body_para(doc, 'Step 2: Dimensionless Groups', bold=True, before=4, after=2)
    equation_box(doc, 'Re = \u03c1\u00b7V\u2097\u00b7d\u1d35/\u03bc = (1.02 \u00d7 5.53 \u00d7 0.04) / 0.012 = 18.8  (laminar)    [Eq. 12]')
    equation_box(doc, 'Sc = \u03bc/(\u03c1\u00b7D) = 0.012 / (1.02 \u00d7 7.0\u202f\u00d7\u202f10\u207b\u2076) = 1681    [Eq. 13]')
    equation_box(doc, 'Gz = Re\u00b7Sc\u00b7(d\u1d35/L) = 18.8 \u00d7 1681 \u00d7 (0.04/30) = 42.10    [Eq. 14]')
    body_para(doc, 'Step 3: Sherwood Number and K\u2097', bold=True, before=4, after=2)
    equation_box(doc, 'Sh = 1.86 \u00d7 Gz\u00b9\u141f\u00b3 = 1.86 \u00d7 (42.10)\u00b9\u141f\u00b3 = 6.47    [Eq. 15]')
    equation_box(doc, 'K\u2097 = Sh \u00d7 D / d\u1d35 = 6.47 \u00d7 7.0\u202f\u00d7\u202f10\u207b\u2076 / 0.04 = 1.1324\u202f\u00d7\u202f10\u207b\u00b3\u202fcm/s    [Eq. 16]')

    body_para(doc,
        'Figure\u202f1 below shows the Sieder-Tate Sh\u2013Gz correlation with the pilot design '
        'point marked. The design point (Gz\u202f=\u202f42.1, Sh\u202f=\u202f6.47) falls on the '
        'theoretical line, confirming validity of the laminar-flow mass transfer model.',
        before=6, after=6)
    p_fig1 = doc.add_paragraph()
    p_fig1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(p_fig1, before=0, after=0)
    run = p_fig1.add_run()
    run.add_picture(fig1_buf, width=Inches(5.2))
    figure_para(doc,
        'Figure\u202f1: Sherwood Number vs Graetz Number. The Sieder-Tate correlation '
        '(Sh\u202f=\u202f1.86\u00b7Gz\u00b9\u141f\u00b3) is shown as the solid blue line. '
        'The BioForge Stage\u202f2 design point (Gz\u202f=\u202f42.1, Sh\u202f=\u202f6.47, '
        'K\u2097\u202f=\u202f1.13\u202f\u00d7\u202f10\u207b\u00b3\u202fcm/s) falls directly '
        'on the theoretical line.')

    section_heading(doc, '5.3  Interfacial Area and Single-Pass Recovery', level=2, before=8)
    equation_box(doc, 'A\u2098\u2092\u2091\u2d35\u2d35\u1d49 = N\u00b7\u03c0\u00b7d\u1d35\u00b7L = 2000\u00b7\u03c0\u00b70.04\u00b730 = 7540\u202fcm\u00b2 (0.7540\u202fm\u00b2)    [Eq. 17]')
    equation_box(doc, '\u03b7\u209b\u1d35\u207f\u1d4f\u2097\u1d49 = 1 \u2212 exp(\u2212K\u2097\u00b7A/Q\u2097) = 1 \u2212 exp(\u22121.1324\u202f\u00d7\u202f10\u207b\u00b3\u00b77540/13.89) = 45.9%    [Eq. 18]')

    section_heading(doc, '5.4  Multi-Module Series Configuration for 95% Recovery', level=2, before=8)
    body_para(doc,
        'A single module achieves 45.9% recovery per pass. For n modules in series:')
    equation_box(doc, '\u03b7\u209c\u2092\u209c\u2090\u2097 = 1 \u2212 (1 \u2212 \u03b7\u209b\u1d35\u207f\u1d4f\u2097\u1d49)\u207f    [Eq. 19]')
    equation_box(doc, 'n = \u2308ln(1\u22120.95) / ln(1\u22120.459)\u2309 = \u23084.87\u2309 = 5 modules    [Eq. 20]')
    equation_box(doc, '\u03b7\u209c\u2092\u209c\u2090\u2097 = 1 \u2212 (0.541)\u2075 = 95.4%    [Eq. 21]')
    body_para(doc,
        'Figure\u202f2 shows cumulative recovery as a function of modules in series. The '
        '5-module configuration is the minimum achieving the 95% target.',
        before=6, after=6)
    p_fig2 = doc.add_paragraph()
    p_fig2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(p_fig2, before=0, after=0)
    run = p_fig2.add_run()
    run.add_picture(fig2_buf, width=Inches(5.2))
    figure_para(doc,
        'Figure\u202f2: Cumulative citric acid recovery vs modules in series. '
        'N\u202f=\u202f2,000 fibers/module, L\u202f=\u202f30\u202fcm, '
        'Q\u2097\u202f=\u202f13.9\u202fcm\u00b3/s. The 5-module design (green bar: 95.4%) '
        'exceeds the 95% target and surpasses the conventional extraction baseline of 85% '
        'starting from module 3.')

    section_heading(doc, '5.5  Pilot-Scale Module Summary', level=2, before=8)
    data_table(doc,
        headers=['Design Output', 'Value', 'Units'],
        rows=[
            ['Liquid velocity in lumen',             '5.53',                   'cm/s'],
            ['Reynolds number',                      '18.8',                   'laminar confirmed (<2,100)'],
            ['Schmidt number',                       '1,681',                  '\u2014'],
            ['Graetz number',                        '42.10',                  '\u2014'],
            ['Sherwood number',                      '6.47',                   '\u2014'],
            ['K\u2097 (liquid-phase mass transfer coefficient)', '1.1324\u202f\u00d7\u202f10\u207b\u00b3', 'cm/s'],
            ['Interfacial area per module',          '7,540 (0.7540\u202fm\u00b2)', 'cm\u00b2'],
            ['Single-pass recovery',                 '45.9%',                  '\u2014'],
            ['Modules for 95% recovery',             '5 (in series)',           '\u2014'],
            ['Total recovery (5 modules)',           '95.4%',                  '\u2014'],
            ['Total membrane area (pilot)',          '3.7699\u202fm\u00b2',    '\u2014'],
        ],
        col_widths=[Inches(3.2), Inches(1.8), Inches(1.8)])

    body_para(doc, 'Figure\u202f3 shows the BioForge Stage\u202f2 pilot-scale module configuration schematic.',
              before=6, after=6)
    p_fig3 = doc.add_paragraph()
    p_fig3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(p_fig3, before=0, after=0)
    run = p_fig3.add_run()
    run.add_picture(fig3_buf, width=Inches(5.5))
    figure_para(doc,
        'Figure\u202f3: BioForge Stage\u202f2 pilot-scale HFC module configuration. Five modules '
        'in series process a 100-L fermentation broth batch at Q\u2097\u202f=\u202f13.9\u202fcm\u00b3/s. '
        'Alamine\u202f336/n-heptane solvent flows on the shell side. Cumulative recovery at each '
        'module outlet is shown. Total area: 3.770\u202fm\u00b2.')
    horizontal_rule(doc)


def section_6(doc):
    section_heading(doc, '6.  Industrial Scale-Up', level=1)
    body_para(doc,
        'Scaling from pilot (100\u202fL) to industrial scale (1,000\u202fL) requires increasing '
        'the liquid flow rate by a factor of 10 while maintaining the same module geometry. '
        'The number of modules in series remains 5, but 10 parallel lanes are added.')
    equation_box(doc, 'Q\u2097,\u1d35\u207f\u1d48 = 1,000,000\u202fcm\u00b3 / (2h\u00b73600\u202fs/h) = 138.9\u202fcm\u00b3/s    [Eq. 22]')
    equation_box(doc, 'n\u209a\u2090\u02b3\u2090\u2097\u2097\u1d49\u2097 = \u230810 parallel lanes\u2309    [Eq. 23]')
    equation_box(doc, 'Total modules = 5 \u00d7 10 = 50     |     Total area = 37.7\u202fm\u00b2    [Eq. 24]')

    data_table(doc,
        headers=['Scale', 'Batch Volume', 'Q\u2097 (cm\u00b3/s)', 'Parallel Lanes',
                 'Total Modules', 'Total Membrane Area', 'Recovery'],
        rows=[
            ['Pilot',      '100\u202fL',    '13.9',  '1',  '5',  '3.770\u202fm\u00b2', '95.4%'],
            ['Industrial', '1,000\u202fL', '138.9', '10', '50', '37.7\u202fm\u00b2',  '95.4%'],
        ],
        col_widths=[Inches(0.8), Inches(1.0), Inches(1.0), Inches(1.0), Inches(1.0), Inches(1.4), Inches(0.9)])

    section_heading(doc, '6.1  Mass Balance: Citric Acid Recovery Comparison', level=2, before=8)
    data_table(doc,
        headers=['Separation Method', 'Recovery', 'CA Recovered (100\u202fL pilot)', 'CA Lost'],
        rows=[
            ['Conventional packed column (baseline)', '85%',   '3,662\u202fg', '646\u202fg'],
            ['BioForge HFC module (this design)',      '95.4%', '4,109\u202fg', '199\u202fg'],
            ['Improvement',                            '+12.2%', '+447\u202fg per batch', '\u2014'],
        ],
        col_widths=[Inches(2.6), Inches(0.9), Inches(1.8), Inches(1.5)])
    horizontal_rule(doc)


def section_7(doc):
    section_heading(doc, '7.  Comparison with Conventional Separation Methods', level=1)
    data_table(doc,
        headers=['Criterion', 'Conventional Packed Column', 'BioForge HFC Module'],
        rows=[
            ['Flooding / channeling',
             'Occurs at high flow rates, limits throughput',
             'Absent: phases separated by membrane'],
            ['Emulsification',
             'Common, causes product contamination',
             'Absent: no direct phase dispersion'],
            ['Scale-up',
             'Requires hydrodynamic redesign at each scale',
             'Linear: add parallel modules, geometry unchanged'],
            ['Footprint',
             '5 to 10\u202fm column height, large base',
             'Compact: 5 modules \u00d7 30\u202fcm each'],
            ['Energy consumption',
             'High (pumping, solvent recovery distillation)',
             'Low: low-pressure pump sufficient'],
            ['Recovery (this system)',
             '85%',
             '95.4%'],
            ['Reproducibility',
             'Depends on packing uniformity, operator skill',
             'Defined by fiber geometry and Q\u2097: fully reproducible'],
            ['Design basis',
             'Empirical correlations with broad uncertainty',
             'Yang-Cussler equations, validated by experiment'],
        ],
        col_widths=[Inches(1.7), Inches(2.4), Inches(2.7)])
    horizontal_rule(doc)


def section_8(doc):
    section_heading(doc, '8.  Research Advances and Engineering Considerations', level=1)
    section_heading(doc, '8.1  Pore Wetting Prevention', level=2, before=8)
    body_para(doc,
        'The most significant engineering risk in liquid-liquid HFC is pore wetting: penetration '
        'of the organic solvent into membrane pores, increasing membrane resistance. The '
        'breakthrough pressure (Young-Laplace equation):')
    equation_box(doc, 'P\u2099\u02b3\u02b3\u02b3\u02b3 = 4\u03b3\u00b7cos(\u03b8) / d\u209a\u2092\u02b3\u1d49,\u2098\u2090\u02e3    [Eq. 25]')
    body_para(doc,
        'For polypropylene (d\u209a\u2092\u02b3\u1d49\u202f=\u202f300\u202f\u00c5\u202f=\u202f'
        '3\u202f\u00d7\u202f10\u207b\u2076\u202fcm), n-heptane surface tension '
        '\u03b3\u202f=\u202f20\u202fmN/m, and contact angle \u03b8\u202f=\u202f0\u00b0, '
        'P\u2099\u02b3\u02b3\u02b3\u02b3\u202f\u2248\u202f2.67\u202fatm. Operating the shell '
        'side at transmembrane pressure below 0.5\u202fatm (standard practice) prevents organic '
        'solvent from entering the pores.')
    section_heading(doc, '8.2  Hybrid Membrane Solution', level=2, before=8)
    body_para(doc,
        'Sharma et al. (2021) demonstrated that dual-layer PVA/PVDF membranes combine a '
        'hydrophobic PVDF top layer (same function as the Yang-Cussler polypropylene membrane) '
        'with a hydrophilic PVA sublayer that prevents total performance loss if the top layer '
        'is compromised over extended operation. BioForge Stage\u202f2 will evaluate hybrid '
        'membranes for systems requiring continuous operation beyond 6-hour batch cycles.')
    section_heading(doc, '8.3  Solvent Selection', level=2, before=8)
    body_para(doc,
        'Alamine\u202f336 dissolved in n-heptane satisfies both design requirements: '
        '(1)\u202fH\u202f=\u202f5 to 15 at pH\u202f2\u20133 (confirming reactive regime), and '
        '(2)\u202fn-heptane surface tension of 20\u202fmN/m against polypropylene is well below '
        'the breakthrough pressure threshold. Back-extraction is achieved by contacting the '
        'loaded organic phase with warm deionized water (60\u202fto\u202f80\u202f\u00b0C) or '
        'dilute sodium hydroxide in a second HFC array \u2014 the same Yang-Cussler equations apply.')
    horizontal_rule(doc)


def section_9(doc):
    section_heading(doc, '9.  BioForge Stage\u202f2 Implementation Plan', level=1)
    section_heading(doc, '9.1  Integration with BioForge Stage\u202f1 Output', level=2, before=8)
    body_para(doc,
        'BioForge Stage\u202f1 produces a protocol card specifying the predicted citric acid yield '
        '(C\u1d35\u207f) for any user-defined waste stream. Stage\u202f2 will accept C\u1d35\u207f '
        'as input and automatically calculate: Q\u2097 for a user-specified batch volume and '
        'processing time, K\u2097 from Eq.\u202f16, required interfacial area A from '
        'Eqs.\u202f17\u201318, number of modules and parallel lanes, and expected recovery. '
        'This produces a complete module specification for separation alongside the fermentation '
        'protocol card.')
    section_heading(doc, '9.2  BioForge-HFC Software Module', level=2, before=8)
    body_para(doc,
        'The Stage\u202f2 separation module will be implemented as '
        'bioforge/separation/hfc_designer.py \u2014 a Python class that accepts Stage\u202f1 '
        'protocol card outputs and computes the full HFC module design. It will output a protocol '
        'card extension listing module count, membrane area, flow configuration, and expected '
        'recovery for any target product from any waste stream.')
    section_heading(doc, '9.3  Planned Experimental Validation', level=2, before=8)
    body_para(doc,
        'The pilot-scale design (5 modules, 3.770\u202fm\u00b2 total area, 95.4% recovery) will '
        'be validated experimentally as part of Stage\u202f1b laboratory work. Protocol: '
        '(1)\u202ffilter fermentation broth through a 0.2-\u03bcm membrane, '
        '(2)\u202fpump at Q\u2097\u202f=\u202f13.9\u202fcm\u00b3/s through the HFC array, '
        '(3)\u202fmeasure inlet/outlet citric acid concentration by HPLC, '
        '(4)\u202fcompare measured K\u2097 to Yang-Cussler prediction '
        '(1.1324\u202f\u00d7\u202f10\u207b\u00b3\u202fcm/s), and '
        '(5)\u202fupdate the BioForge Stage\u202f2 software module accordingly.')

    callout_box(doc,
        'Stage\u202f2 Module: What BioForge Will Produce',
        [
            'When a U.S. researcher or manufacturer runs their waste stream through BioForge, '
            'they will receive two documents: (1)\u202fa fermentation optimization protocol card '
            '(Stage\u202f1) specifying optimal process stimulants and predicted yield, and '
            '(2)\u202fa separation module specification (Stage\u202f2) sizing the hollow-fiber '
            'extraction array required to recover that product. Together, these constitute a '
            'complete, independently reproducible waste-to-value production protocol.'
        ],
        bg_color=INFO_BG, border_color=INFO_BDR)
    horizontal_rule(doc)


def section_10(doc):
    section_heading(doc, '10.  Conclusion', level=1)
    body_para(doc,
        'This study applies the Yang-Cussler (1986) hollow-fiber membrane contactor design '
        'framework to a specific, practically relevant separation problem: recovery of citric '
        'acid from yam peel agricultural waste fermentation broth at the yield validated by my '
        'published research (Okedi et al., 2024) and reproduced computationally by BioForge '
        'Stage\u202f1.')
    body_para(doc, 'Key design results:', bold=True, before=6, after=4)
    for item in [
        'Liquid-phase mass transfer coefficient: K\u2097\u202f=\u202f1.1324\u202f\u00d7\u202f'
        '10\u207b\u00b3\u202fcm/s at pilot-scale operating conditions (Re\u202f=\u202f18.8, laminar).',
        'Single-pass recovery: 45.9% per module (N\u202f=\u202f2,000 fibers, L\u202f=\u202f30\u202fcm, '
        'A\u202f=\u202f7,540\u202fcm\u00b2).',
        'Modules in series for 95% recovery: 5, achieving 95.4% cumulative recovery.',
        'Total pilot membrane area: 3.770\u202fm\u00b2 for a 100-L batch in 2 hours.',
        'Industrial scale (1,000\u202fL): 50 modules (5 series \u00d7 10 parallel), '
        '37.7\u202fm\u00b2 total area.',
        'Recovery improvement over conventional extraction (85%): +12.2%, yielding 447 '
        'additional grams of citric acid per 100-L batch.',
    ]:
        bullet_para(doc, item)
    body_para(doc,
        'The Yang-Cussler Sieder-Tate correlation applies directly when Re\u202f=\u202f18.8, '
        'confirming laminar flow. Large H for Alamine\u202f336 reduces the three-resistance '
        'equation to K\u202f\u2248\u202fK\u2097, and the polypropylene fiber geometry is identical '
        'to the Yang-Cussler experimental setup.',
        before=6)
    body_para(doc,
        'This report demonstrates that the petitioner has the theoretical knowledge, quantitative '
        'modeling skill, and domain-specific engineering judgment required to design and implement '
        'BioForge Stage\u202f2. The hollow-fiber separation module designed here, integrated with '
        'the BioForge Stage\u202f1 optimization engine, will produce a complete waste-to-value '
        'process design that any U.S. laboratory or manufacturer can reproduce independently.')
    horizontal_rule(doc)


def nomenclature(doc):
    section_heading(doc, 'Nomenclature', level=1)
    data_table(doc,
        headers=['Symbol', 'Definition', 'Value / Units'],
        rows=[
            ['A',    'Interfacial area of hollow-fiber module',          'cm\u00b2'],
            ['D',    'Citric acid diffusivity in water at 30\u202f\u00b0C', '7.0\u202f\u00d7\u202f10\u207b\u2076\u202fcm\u00b2/s'],
            ['d\u1d35', 'Fiber inner diameter',                          '0.04\u202fcm'],
            ['d\u2092', 'Fiber outer diameter',                          '0.046\u202fcm'],
            ['Gz',   'Graetz number = Re\u00b7Sc\u00b7(d\u1d35/L)',      '42.10 (this design)'],
            ['H',    'Distribution coefficient (organic/aqueous)',        '5\u201315 (Alamine\u202f336 / citric acid)'],
            ['K',    'Overall mass transfer coefficient',                 'cm/s'],
            ['K\u2097', 'Liquid-phase mass transfer coefficient',         '1.1324\u202f\u00d7\u202f10\u207b\u00b3\u202fcm/s (this design)'],
            ['L',    'Module length',                                     '30\u202fcm'],
            ['\u03bc', 'Broth viscosity',                                 '0.012\u202fg/(cm\u00b7s)'],
            ['N',    'Number of fibers per module',                       '2,000'],
            ['Pe',   'Peclet number = Re\u00b7Sc',                        '31,578'],
            ['Q\u2097', 'Liquid volumetric flow rate',                    '13.89\u202fcm\u00b3/s (pilot)'],
            ['Re',   'Reynolds number = \u03c1\u00b7V\u2097\u00b7d\u1d35/\u03bc', '18.8 (laminar)'],
            ['\u03c1', 'Broth density',                                   '1.02\u202fg/cm\u00b3'],
            ['Sc',   'Schmidt number = \u03bc/(\u03c1\u00b7D)',           '1,681'],
            ['Sh',   'Sherwood number = K\u2097\u00b7d\u1d35/D',         '6.47 (this design)'],
            ['t\u1d42', 'Fiber wall thickness',                           '0.003\u202fcm'],
            ['V\u2097', 'Liquid velocity in fiber lumen',                 '5.53\u202fcm/s'],
            ['\u03b5', 'Membrane porosity',                               '0.33 (33%)'],
            ['\u03b7', 'Fractional recovery',                             '0.954 (95.4% at 5 modules)'],
        ],
        col_widths=[Inches(0.8), Inches(3.5), Inches(2.5)])
    horizontal_rule(doc)


def appendix_a(doc):
    section_heading(doc, 'Appendix A: Complete Calculation Summary', level=1)
    data_table(doc,
        headers=['Step', 'Quantity', 'Equation', 'Result'],
        rows=[
            ['1',  'A\u2091\u02b3\u02b3\u02b3\u02b3', 'N\u00b7\u03c0\u00b7(d\u1d35/2)\u00b2', '2.5133\u202fcm\u00b2'],
            ['2',  'V\u2097',         'Q\u2097 / A\u2091\u02b3',           '5.53\u202fcm/s'],
            ['3',  'Re',              '\u03c1\u00b7V\u2097\u00b7d\u1d35 / \u03bc', '18.8 (laminar)'],
            ['4',  'Sc',              '\u03bc / (\u03c1\u00b7D)',            '1,681'],
            ['5',  'Gz',              'Re\u00b7Sc\u00b7(d\u1d35/L)',        '42.10'],
            ['6',  'Sh',              '1.86 \u00b7 Gz\u00b9\u141f\u00b3',   '6.47'],
            ['7',  'K\u2097',         'Sh \u00d7 D / d\u1d35',             '1.1324\u202f\u00d7\u202f10\u207b\u00b3\u202fcm/s'],
            ['8',  'A\u2098\u2092\u1d49', 'N\u00b7\u03c0\u00b7d\u1d35\u00b7L', '7,540\u202fcm\u00b2 (0.7540\u202fm\u00b2)'],
            ['9',  '\u03b7\u209b\u1d35\u207f\u1d4f\u2097\u1d49', '1 \u2212 exp(\u2212K\u2097\u00b7A/Q\u2097)', '45.9%'],
            ['10', 'n\u209b\u1d49\u02b3\u1d49\u1d49', '\u2308ln(0.05)/ln(1\u2212\u03b7\u209b\u1d35\u207f\u1d4f\u2097\u1d49)\u2309', '5 modules'],
            ['11', '\u03b7\u209c\u2092\u209c\u2090\u2097', '1 \u2212 (0.541)\u2075', '95.4%'],
        ],
        col_widths=[Inches(0.5), Inches(1.0), Inches(3.2), Inches(2.0)])
    horizontal_rule(doc)


def appendix_b(doc):
    section_heading(doc, 'Appendix B: MATLAB Design Code', level=1)
    body_para(doc,
        'The following MATLAB script reproduces all calculations and serves as a validation '
        'tool or starting point for parameter sensitivity studies.')
    code_lines = [
        '% BioForge Stage 2 -- Hollow Fiber Contactor Design',
        '% Citric Acid Recovery from Yam Peel Fermentation Broth',
        '% Ogaga Maxwell Okedi, March 2026',
        'clear; clc;',
        '',
        '% Process specifications',
        'C_feed     = 43.08;             % g/L citric acid (Okedi et al., 2024)',
        'eta_target = 0.95;              % Target recovery (95%)',
        'Q_L        = 100e3 / (2*3600);  % cm3/s (100 L batch, 2 h)',
        '',
        '% Physical properties at 30 deg C',
        'D   = 7.0e-6;   % cm^2/s, citric acid diffusivity in water',
        'mu  = 0.012;    % g/(cm*s), broth viscosity',
        'rho = 1.02;     % g/cm^3, broth density',
        '',
        '% Fiber and module specs (Yang-Cussler geometry)',
        'd_i = 0.04;  d_o = 0.046;  t_w = 0.003;',
        'N = 2000;  L = 30;  % fibers, cm',
        '',
        '% Step 1: Cross-sectional area and velocity',
        'A_cross = N * pi * (d_i/2)^2;',
        'V_L     = Q_L / A_cross;',
        '',
        '% Step 2: Dimensionless groups',
        'Re = rho * V_L * d_i / mu;',
        'Sc = mu / (rho * D);',
        'Gz = Re * Sc * (d_i / L);',
        '',
        '% Step 3: Sherwood number and KL',
        'Sh  = 1.86 * Gz^(1/3);',
        'K_L = Sh * D / d_i;',
        '',
        '% Step 4: Area, recovery, modules in series',
        'A_module   = N * pi * d_i * L;',
        'eta_single = 1 - exp(-K_L * A_module / Q_L);',
        'n_series   = ceil(log(1-eta_target) / log(1-eta_single));',
        'eta_total  = 1 - (1-eta_single)^n_series;',
        '',
        '% Industrial scale-up',
        'n_parallel = ceil((1000e3/(2*3600)) / Q_L);',
        'total_mods = n_series * n_parallel;',
        'total_area = total_mods * A_module / 1e4;  % m^2',
        '',
        '% Display results',
        'fprintf("KL = %.4e cm/s\\n", K_L);',
        'fprintf("Modules for 95%%: %d  (eta = %.1f%%)\\n", n_series, eta_total*100);',
        'fprintf("Industrial: %d modules, %.1f m^2\\n", total_mods, total_area);',
    ]
    tbl  = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    _set_cell_color(cell, EQ_BG)
    _no_borders(cell)
    _cell_margins(cell, top=80, bottom=80, left=160, right=80)
    first = True
    for line in code_lines:
        if first:
            p = cell.paragraphs[0]; first = False
        else:
            p = cell.add_paragraph()
        _para_spacing(p, before=0, after=0, line=1.05)
        _run(p, line if line else ' ', size=9, color=DARK_GRAY, mono=True)
    doc.add_paragraph()
    horizontal_rule(doc)


def references(doc):
    section_heading(doc, 'References', level=1)
    refs = [
        'Yang, M.; Cussler, E.\u202fL. Designing Hollow-Fiber Contactors. AIChE J. 1986, 32 (11), 1910\u20131916.',
        'Okedi, O.\u202fM. et al. Optimization of Citric Acid Production from Yam Peel Agricultural Waste Using '
        'Box-Behnken Design and ANFIS. Industrial Crops and Products 2024 (IF\u202f6.2).',
        'Paul, A.\u202fR.; Callahan, R.\u202fW. Liquid-Liquid Extraction and Stripping of Gold with Microporous '
        'Hollow Fibers. J. Membr. Sci. 1987, 35 (1), 57\u201371.',
        'Cussler, E.\u202fL. Diffusion: Mass Transfer in Fluid Systems; Cambridge University Press: Cambridge, 1984.',
        'Qi, Z.; Cussler, E.\u202fL. Microporous Hollow Fibers for Gas Absorption I: Mass Transfer in the Liquid. '
        'J. Membr. Sci. 1985.',
        'Gableman, A.; Hwang, S.-T. Hollow Fiber Membrane Contactors. J. Membr. Sci. 1999.',
        'Sharma, A.\u202fK. et al. Porous Hydrophobic-Hydrophilic Composite Hollow Fiber Membranes for DCMD. '
        'Membranes 2021, 11 (2), 120.',
        'Schmitt, A. et al. Evaluation of an Ozone Diffusion Process Using a Hollow Fiber Membrane Contactor. '
        'Chem. Eng. Res. Des. 2022, 177, 291\u2013303.',
        'Bazhenov, S. et al. Gas-Liquid Hollow Fiber Membrane Contactors for Different Applications. Fibers 2018, '
        '6 (4), 76.',
        'Ramezani, R. et al. A Review on Hollow Fiber Membrane Contactors for Carbon Capture. Processes 2022, '
        '10 (10), 2103.',
        'Li, K.; Tai, M.\u202fS.\u202fL.; Teo, W.\u202fK. Design of a CO\u2082 Scrubber Using a Microporous '
        'Membrane. J. Membr. Sci. 1994, 86, 119\u2013125.',
        'Tai, M.\u202fS.\u202fL. et al. Removal of Dissolved Oxygen in Ultrapure Water Using Microporous Membrane '
        'Modules. J. Membr. Sci. 1994, 87, 99\u2013105.',
        'Kneifel, K. et al. Hollow Fiber Membrane Contactor for Air Humidity Control. J. Membr. Sci. 2006, 276, '
        '241\u2013251.',
        'Sieder, E.\u202fN.; Tate, G.\u202fE. Heat Transfer and Pressure Drop of Liquids in Tubes. Ind. Eng. Chem. '
        '1936, 28 (12), 1429\u20131435.',
        'Boucif, N. et al. Hollow Fiber Membrane Contactor for Hydrogen Sulfide Odor Control. AIChE J. 2008, 54 (1), '
        '122\u2013131.',
    ]
    for i, ref in enumerate(refs, start=1):
        p = doc.add_paragraph()
        _para_spacing(p, before=2, after=2, line=1.1)
        p.paragraph_format.left_indent       = Inches(0.35)
        p.paragraph_format.first_line_indent = Inches(-0.35)
        _run(p, f'{i}.  ', bold=True, size=10, color=DARK_GRAY)
        _run(p, ref, size=10, color=DARK_GRAY)


# ══════════════════════════════════════════════════════════════════════════════
#  MAIN
# ══════════════════════════════════════════════════════════════════════════════

def main():
    print('Generating figures...')
    fig1 = _fig_sh_gz()
    fig2 = _fig_recovery()
    fig3 = _fig_schematic()

    print('Building document...')
    doc = Document()
    _set_page_margins(doc)

    # Remove default paragraph spacing from Normal style
    doc.styles['Normal'].paragraph_format.space_before = Pt(0)
    doc.styles['Normal'].paragraph_format.space_after  = Pt(0)
    doc.styles['Normal'].font.name = BODY_FONT
    doc.styles['Normal'].font.size = Pt(11)

    title_block(doc)
    section_1(doc)
    section_2(doc)
    section_3(doc)
    section_4(doc)
    section_5(doc, fig1, fig2, fig3)
    section_6(doc)
    section_7(doc)
    section_8(doc)
    section_9(doc)
    section_10(doc)
    nomenclature(doc)
    appendix_a(doc)
    appendix_b(doc)
    references(doc)

    doc.save(OUT)
    size_kb = os.path.getsize(OUT) // 1024
    print(f'Written: {OUT}')
    print(f'Size: {size_kb} KB')


if __name__ == '__main__':
    main()
